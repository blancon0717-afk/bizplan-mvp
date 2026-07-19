"""생성된 섹션들을 DOCX로 내보내기.

출력 규칙:
- 좌측 사업계획 본문 내용만 출력 (메모/검토영역/신뢰도 배지 미포함)
- user_answer 세그먼트: 기본 검정
- llm_inferred 세그먼트: 회색 (RGB 128,128,128)
- [출처 필요] / [추정값] / [수치 필요] 태그: 빨간색
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from core.forms import Form
from core.generation import ContentSegment, SectionResult

_FONT = "맑은 고딕"
_COLOR_BLACK = RGBColor(17, 17, 17)
_COLOR_BLUE = RGBColor(59, 130, 246)
_COLOR_BLUE_DARK = RGBColor(29, 78, 216)
_MD_TABLE_ROW = re.compile(r"^\s*\|")
_BULLET_LINE = re.compile(r"^\s*-\s")
_HEADING_LINE = re.compile(r"^■")
_CAPTION_CELL = re.compile(r"^<.*>$")
_SOURCE_CELL = re.compile(r"^출처:")
_DESCRIPTION_CELL = re.compile(r"^\[")
_URL_PATTERN = re.compile(r"https?://\S+")


def _set_font(run, size_pt: float, bold: bool, color: RGBColor) -> None:
    run.font.name = _FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), _FONT)
    rfonts.set(qn("w:hAnsi"), _FONT)
    rfonts.set(qn("w:eastAsia"), _FONT)


def _set_cell_bg(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _add_runs(para, text: str, size_pt: float, bold: bool, base_color: RGBColor) -> None:
    if not text:
        return
    run = para.add_run(text)
    _set_font(run, size_pt, bold, base_color)


def _set_table_borders(tbl) -> None:
    """모든 변에 단선 테두리 지정 — 'Table Grid' 스타일이 없는 문서용."""
    tbl_pr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def _parse_md_table(lines: list[str]) -> tuple[list[list[str]], list[list[str]]]:
    def parse_row(line: str) -> list[str]:
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        return cells

    rows = [parse_row(l) for l in lines]
    sep_idx = next(
        (i for i, r in enumerate(rows) if r and all(re.match(r"^[-:\s]+$", c) for c in r)),
        -1,
    )
    headers = rows[:sep_idx] if sep_idx >= 0 else []
    body = rows[sep_idx + 1:] if sep_idx >= 0 else rows
    return headers, body


def _add_segment(doc: Document, seg: ContentSegment) -> None:
    base_color = _COLOR_BLACK
    lines = seg.text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 마크다운 표
        if _MD_TABLE_ROW.match(line):
            table_block: list[str] = []
            while i < len(lines) and _MD_TABLE_ROW.match(lines[i]):
                table_block.append(lines[i])
                i += 1
            headers, body = _parse_md_table(table_block)
            all_rows = headers + body
            if not all_rows:
                continue
            num_cols = max(len(r) for r in all_rows)
            tbl = doc.add_table(rows=len(all_rows), cols=num_cols)
            try:
                tbl.style = "Table Grid"
            except KeyError:
                # 공식 원본 템플릿(한컴 변환)에는 'Table Grid' 스타일이 없음 → 테두리 직접 지정
                _set_table_borders(tbl)
            for ri, row_cells in enumerate(all_rows):
                is_header = ri < len(headers)
                is_caption = not is_header and bool(row_cells) and all(
                    ct.strip() == "" or bool(_CAPTION_CELL.match(ct.strip()))
                    for ct in row_cells
                )
                for ci in range(num_cols):
                    cell = tbl.cell(ri, ci)
                    ct = (row_cells[ci] if ci < len(row_cells) else "").strip()
                    para = cell.paragraphs[0]
                    if is_header:
                        _set_cell_bg(cell, "F8FAFC")
                        _add_runs(para, ct, 9, True, base_color)
                    elif is_caption:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run(ct)
                        _set_font(run, 8, False, _COLOR_BLACK)
                        run.font.italic = True
                    elif _DESCRIPTION_CELL.match(ct):
                        _set_cell_bg(cell, "DBEAFE")
                        run = para.add_run(ct)
                        _set_font(run, 9, False, _COLOR_BLUE_DARK)
                        run.font.italic = True
                    elif _SOURCE_CELL.match(ct):
                        url_m = _URL_PATTERN.search(ct)
                        if url_m:
                            run_pre = para.add_run(ct[:url_m.start()])
                            _set_font(run_pre, 8, False, _COLOR_BLACK)
                            run_url = para.add_run(url_m.group())
                            _set_font(run_url, 8, False, _COLOR_BLUE)
                            run_url.font.underline = True
                            tail = ct[url_m.end():]
                            if tail:
                                run_tail = para.add_run(tail)
                                _set_font(run_tail, 8, False, _COLOR_BLACK)
                        else:
                            run = para.add_run(ct)
                            _set_font(run, 8, False, _COLOR_BLUE)
                    else:
                        _add_runs(para, ct, 9, False, base_color)
            doc.add_paragraph()

        # 빈 줄
        elif line.strip() == "":
            doc.add_paragraph()
            i += 1

        # ■ 중제목
        elif _HEADING_LINE.match(line):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(0)
            _add_runs(para, line, 10, True, base_color)
            i += 1

        # 세부항목 ( - )
        elif _BULLET_LINE.match(line):
            bullet_lines: list[str] = []
            while i < len(lines) and _BULLET_LINE.match(lines[i]):
                bullet_lines.append(lines[i])
                i += 1
            for bl in bullet_lines:
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.left_indent = Cm(0.5)
                _add_runs(para, bl, 10, False, base_color)

        # 일반 단락
        else:
            text_lines: list[str] = []
            while (
                i < len(lines)
                and not _MD_TABLE_ROW.match(lines[i])
                and lines[i].strip() != ""
                and not _HEADING_LINE.match(lines[i])
                and not _BULLET_LINE.match(lines[i])
            ):
                text_lines.append(lines[i])
                i += 1
            joined = "\n".join(text_lines).strip()
            if joined:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(4)
                _add_runs(para, joined, 10, False, base_color)


def export_to_docx(
    form: Form | None,
    results: list[SectionResult],
    business_name: str = "(미지정)",
    title: str | None = None,
) -> BytesIO:
    """섹션 결과를 DOCX로 변환.

    title이 주어지면 표지 제목으로 사용(양식 없는 초안 export용).
    없으면 form.program_name을 사용한다.
    """
    doc = Document()

    # 여백 설정 (25mm)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Normal 스타일 기본 폰트를 맑은 고딕으로 (ASCII + 한글 모두)
    style = doc.styles["Normal"]
    style.font.name = _FONT  # ascii, hAnsi 설정
    style.font.size = Pt(10)
    # style.font.name이 rPr/rFonts를 생성한 뒤 eastAsia 추가
    style_rpr = style.element.find(qn("w:rPr"))
    if style_rpr is not None:
        style_rfonts = style_rpr.find(qn("w:rFonts"))
        if style_rfonts is not None:
            style_rfonts.set(qn("w:eastAsia"), _FONT)

    # 표지 제목
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    program_name = title or (form.program_name if form else "사업계획서")
    run = title_para.add_run(program_name)
    _set_font(run, 16, True, _COLOR_BLACK)

    # 기업명
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"기업명: {business_name}")
    _set_font(run, 10, False, _COLOR_BLACK)

    doc.add_paragraph()  # 제목 아래 여백

    for r in results:
        # 섹션 대제목
        heading_para = doc.add_paragraph()
        heading_para.paragraph_format.space_before = Pt(8)
        heading_para.paragraph_format.space_after = Pt(4)
        run = heading_para.add_run(f"[{r.section_id}] {r.section_title}")
        _set_font(run, 12, True, _COLOR_BLACK)

        if r.user_edited_content is not None:
            _add_segment(doc, ContentSegment(text=r.user_edited_content, source="user_answer"))
        else:
            segments = r.content_segments or [
                ContentSegment(text=r.content or "", source="llm_inferred")
            ]
            for seg in segments:
                if seg.text.strip():
                    _add_segment(doc, seg)

        doc.add_paragraph()  # 섹션 간 여백

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ────────────────────────────────────────────────────────────────────
# 공식 양식 템플릿 렌더링 (docxtpl)
#
# data/templates/{program_code}.docx — 공식 양식 레이아웃을 재현한 템플릿에
# 변환 결과를 서브독으로 주입한다. 템플릿·의존성이 없으면 None을 반환해
# 호출측이 기존 일반 렌더러(export_to_docx)로 폴백한다.
# ────────────────────────────────────────────────────────────────────

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "data" / "templates"

# 템플릿이 {{p sections[...] }} 서브독으로 받는 본문 섹션 id
# (scripts/prepare_official_templates.py 가 심는 태그와 반드시 일치)
TEMPLATE_SUBDOC_IDS: dict[str, tuple[str, ...]] = {
    "deeptech_academy": ("1-1", "1-2", "1-3", "1-4", "2-1", "2-2", "2-3", "2-4", "2-5",
                         "3-1", "3-2", "3-3", "3-4", "4-1", "4-2"),
    "initial_package": ("1", "2", "3", "4"),
}

# 정적 표의 과제명/아이템명 셀에 넣을 제안값 추출용
_PROJECT_NAME_RE = re.compile(
    r"\|\s*(?:사업화 과제명|창업아이템명)[^|]*\|\s*([^|\n]+?)\s*\|"
)


def _extract_project_name(results: list[SectionResult]) -> str:
    """0-x 섹션 변환 결과에서 과제명/창업아이템명 제안값 추출. 없으면 빈 문자열."""
    for r in results:
        if not str(r.section_id).startswith("0"):
            continue
        text = r.user_edited_content if r.user_edited_content is not None else (r.content or "")
        m = _PROJECT_NAME_RE.search(text)
        if not m:
            continue
        val = m.group(1).strip().strip("()").strip()
        # 가이드 문구가 그대로 남은 경우 제외
        if not val or "제안" in val or "초안" in val or "기재" in val:
            continue
        return val[:80]
    return ""


# ── markdown 표 파서 (LLM 출력 → 템플릿 셀 값) ──────────────────────

def _md_tables(text: str) -> list[list[list[str]]]:
    """markdown 표들을 [표][행][셀] 리스트로 추출 (구분선 행 제외)."""
    tables: list[list[list[str]]] = []
    cur: list[list[str]] = []
    for line in (text or "").split("\n"):
        if _MD_TABLE_ROW.match(line):
            cells = [c.strip() for c in line.split("|")]
            if cells and cells[0] == "":
                cells = cells[1:]
            if cells and cells[-1] == "":
                cells = cells[:-1]
            if cells and not all(re.match(r"^[-:\s]*$", c) for c in cells):
                cur.append(cells)
        elif cur:
            tables.append(cur)
            cur = []
    if cur:
        tables.append(cur)
    return tables


def _strip_md_tables(text: str) -> str:
    return "\n".join(l for l in (text or "").split("\n") if not _MD_TABLE_ROW.match(l))


def _kv_get(kv: dict[str, str], key: str) -> str:
    """라벨 매칭 — 정확 일치 우선, 다음 괄호 제거·전방 일치."""
    if key in kv:
        return kv[key]
    nk = key.split("(")[0].strip()
    for k, v in kv.items():
        kk = k.split("(")[0].strip()
        if kk.startswith(nk) or nk.startswith(kk):
            return v
    return ""


def _section_text(results_by_id: dict, sid: str) -> str:
    r = results_by_id.get(sid)
    if r is None:
        return ""
    return r.user_edited_content if r.user_edited_content is not None else (r.content or "")


def _subdoc_from_text(tpl, text: str):
    sd = tpl.new_subdoc()
    if text.strip():
        _add_segment(sd, ContentSegment(text=text, source="llm_inferred"))
    return sd


def _subdoc_from_result(tpl, r: SectionResult | None):
    sd = tpl.new_subdoc()
    if r is None:
        return sd
    if r.user_edited_content is not None:
        segs = [ContentSegment(text=r.user_edited_content, source="user_answer")]
    else:
        segs = r.content_segments or [ContentSegment(text=r.content or "", source="llm_inferred")]
    for seg in segs:
        if seg.text.strip():
            _add_segment(sd, seg)
    return sd


# ── 폼별 컨텍스트 빌더 ────────────────────────────────────────────────

def _sections_ctx(tpl, by_id: dict, ids: tuple[str, ...]) -> dict:
    return {sid: _subdoc_from_result(tpl, by_id.get(sid)) for sid in ids}


def _first_kv(text: str) -> dict[str, str]:
    """첫 번째 2열 이상 markdown 표 → {열0: 열1} (헤더 행 제외)."""
    for tbl in _md_tables(text):
        kv = {}
        for row in tbl:
            if len(row) >= 2 and row[0] and row[0] not in ("항목", "분야", "구분"):
                kv[row[0]] = row[1]
        if kv:
            return kv
    return {}


def _ctx_deeptech(tpl, by_id: dict) -> dict:
    kv = _first_kv(_section_text(by_id, "0-3"))
    ov_keys = ("명칭", "범주", "회사 사이트", "소개", "진출 목표시장",
               "경쟁사 대비 차별성", "현황 및 구체화 방안")
    return {
        "sections": _sections_ctx(tpl, by_id, TEMPLATE_SUBDOC_IDS["deeptech_academy"]),
        "ov": {k: _kv_get(kv, k) for k in ov_keys},
    }


def _ctx_initial(tpl, by_id: dict) -> dict:
    kv = _first_kv(_section_text(by_id, "0-2"))
    ov_keys = ("명칭", "범주", "아이템 개요", "문제 인식", "실현 가능성", "성장전략", "팀 구성")
    return {
        "sections": _sections_ctx(tpl, by_id, TEMPLATE_SUBDOC_IDS["initial_package"]),
        "ov": {k: _kv_get(kv, k) for k in ov_keys},
    }


def _ctx_voucher(tpl, by_id: dict) -> dict:
    kv2 = _first_kv(_section_text(by_id, "2"))
    kv4 = _first_kv(_section_text(by_id, "4"))
    # 간트: '월' 헤더 표에서 {분야: {월숫자: 셀}}
    g = {f: {str(m): "" for m in range(2, 11)} for f in ("컨설팅", "기술지원", "마케팅")}
    for tbl in _md_tables(_section_text(by_id, "5")):
        header = tbl[0]
        month_cols = {ci: re.sub(r"\D", "", h) for ci, h in enumerate(header) if re.search(r"\d", h)}
        if not month_cols:
            continue
        for row in tbl[1:]:
            field = row[0].strip()
            if field in g:
                for ci, m in month_cols.items():
                    if ci < len(row) and m in g[field]:
                        g[field][m] = row[ci].strip()
        break
    # KPI: 3열 표(구분|지표|값) → {지표: 값}
    kpi_src: dict[str, str] = {}
    for tbl in _md_tables(_section_text(by_id, "6")):
        for row in tbl:
            if len(row) >= 3 and row[1] and row[1] != "지표":
                kpi_src[row[1]] = row[2]
        if kpi_src:
            break
    kpi_keys = ("고용 증가율(%)", "신규 고용인원수(명)", "매출액 증가율(%)", "매출 증가액(천원)")
    return {
        "kv2": {k: _kv_get(kv2, k) for k in
                ("제품용도 및 특성", "제품생산 공정", "시장 상황", "기술품질 경쟁력", "지식재산권 및 인증 보유현황")},
        "kv4": {k: _kv_get(kv4, k) for k in ("컨설팅", "기술지원", "마케팅")},
        "g": g,
        "kpi": {k: _kv_get(kpi_src, k) for k in kpi_keys},
        "sec3": _subdoc_from_text(tpl, _section_text(by_id, "3")),
        "sec6": _subdoc_from_text(tpl, _strip_md_tables(_section_text(by_id, "6"))),
    }


_CTX_BUILDERS = {
    "deeptech_academy": _ctx_deeptech,
    "initial_package": _ctx_initial,
    "innovation_voucher": _ctx_voucher,
}


def export_to_official_docx(
    form: Form,
    results: list[SectionResult],
    business_name: str = "(미지정)",
) -> BytesIO | None:
    """공식 양식 원본 템플릿에 변환 결과를 채워 DOCX 생성.

    템플릿 파일이 없거나 docxtpl 미설치·렌더 실패 시 None 반환(호출측 폴백).
    """
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        logger.warning("[DOCX] docxtpl 미설치 — 일반 렌더러로 폴백")
        return None

    tpl_path = _TEMPLATES_DIR / f"{form.program_code}.docx"
    builder = _CTX_BUILDERS.get(form.program_code)
    if not tpl_path.exists() or builder is None:
        return None

    try:
        tpl = DocxTemplate(str(tpl_path))
        by_id = {r.section_id: r for r in results}
        ctx = builder(tpl, by_id)
        ctx.setdefault("business_name", business_name)
        ctx.setdefault("project_name", _extract_project_name(results))
        tpl.render(ctx)
        buf = BytesIO()
        tpl.save(buf)
        buf.seek(0)
        return buf
    except Exception as e:  # noqa: BLE001 — 템플릿 문제로 다운로드 자체가 막히면 안 됨
        logger.error("[DOCX] 공식 템플릿 렌더 실패(%s): %s — 일반 렌더러 폴백", form.program_code, e)
        return None
