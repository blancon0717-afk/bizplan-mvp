"""공식 사업계획서 DOCX 템플릿 생성기 (docxtpl용).

공식 양식 PDF의 목차·표 레이아웃을 재현한 템플릿 3종을 생성한다:
  data/templates/deeptech_academy.docx
  data/templates/initial_package.docx
  data/templates/innovation_voucher.docx

템플릿 안의 Jinja 태그:
  {{ business_name }}          — 표지 기업명
  {{ project_name }}           — 과제명/창업아이템명 셀 (변환 결과에서 추출)
  {{p sections["<id>"] }}      — 동적 섹션 본문(서브독) 주입 지점

정적(스켈레톤) 섹션 — 공식 빈 표를 템플릿에 직접 내장(LLM 콘텐츠 미주입):
  deeptech_academy: 0-1, 0-2 / initial_package: 0-1 / innovation_voucher: 1
(core/docx_export.py의 TEMPLATE_STATIC_SECTIONS와 반드시 일치시킬 것)

실행: python scripts/build_form_templates.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from core.forms import load_form  # noqa: E402

OUT_DIR = ROOT / "data" / "templates"
FONT = "맑은 고딕"
BLACK = RGBColor(0x00, 0x00, 0x00)

# 정적 섹션(공식 빈 표 내장) — docx_export.TEMPLATE_STATIC_SECTIONS와 동일해야 함
STATIC_SECTIONS: dict[str, set[str]] = {
    "deeptech_academy": {"0-1", "0-2"},
    "initial_package": {"0-1"},
    "innovation_voucher": {"1"},
}


def _set_font(run, size_pt: float, bold: bool) -> None:
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def _shade(cell, fill_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _cell_text(cell, text: str, size: float = 9, bold: bool = False,
               center: bool = False) -> None:
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _set_font(run, size, bold)


def _base_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    rpr = style.element.find(qn("w:rPr"))
    if rpr is not None:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            rfonts.set(qn("w:eastAsia"), FONT)
    return doc


def _title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    _set_font(run, 16, True)


def _subtitle(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _set_font(run, 10, False)


def _heading(doc: Document, text: str, size: float = 12) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    _set_font(run, size, True)


def _note(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    _set_font(run, 8.5, False)


def _dyn(doc: Document, sec_id: str) -> None:
    """동적 섹션 서브독 주입 태그."""
    doc.add_paragraph(f'{{{{p sections["{sec_id}"] }}}}')


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """항목/내용 2열 표 — 좌측 라벨 음영."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        _shade(tbl.cell(i, 0), "E7E6E6")
        _cell_text(tbl.cell(i, 0), label, bold=True, center=True)
        _cell_text(tbl.cell(i, 1), value)
    doc.add_paragraph()


def _grid_table(doc: Document, header: list[str], blank_rows: int = 2,
                body_rows: list[list[str]] | None = None) -> None:
    """헤더 음영 + 본문 행 표."""
    body_rows = body_rows or []
    n = 1 + len(body_rows) + blank_rows
    tbl = doc.add_table(rows=n, cols=len(header))
    tbl.style = "Table Grid"
    for c, h in enumerate(header):
        _shade(tbl.cell(0, c), "E7E6E6")
        _cell_text(tbl.cell(0, c), h, bold=True, center=True)
    for r, row in enumerate(body_rows, start=1):
        for c, v in enumerate(row):
            if c < len(header):
                _cell_text(tbl.cell(r, c), v)
    doc.add_paragraph()


def _section_label(sec_id: str, title: str) -> str:
    """공식 목차 표기: 본문 섹션은 'id. 제목', 0-x 는 제목만."""
    if re.match(r"^0-", sec_id):
        return f"□ {title}"
    return f"{sec_id}. {title}"


# ────────────────────────────────────────────────────────────────────
# 폼별 정적(스켈레톤) 섹션 — 공식 빈 표
# ────────────────────────────────────────────────────────────────────

def _static_deeptech_0_1(doc: Document) -> None:
    _heading(doc, "1. 신청자 기본 정보", 13)
    _kv_table(doc, [
        ("권역(택1)", "□ 수도권   □ 호남권   □ 영남권"),
        ("사업화 과제명", "{{ project_name }}"),
        ("신청자 성명(생년월일)", ""),
        ("성별", "□ 남  /  □ 여"),
        ("자택주소", ""),
        ("기업명", ""),
        ("사업자등록번호", ""),
        ("법인등록번호(해당시)", ""),
        ("사업장 주소", ""),
        ("사업개시일(회사성립연월일)", ""),
        ("사업자 구분", "□ 개인사업자 □ 법인사업자  /  □ 단독 □ 공동 □ 각자대표"),
    ])
    _note(doc, "※ KPI 설정 (필수 / 계량·비계량 각 1개 선택)")
    _grid_table(doc, ["구분", "선택지"], blank_rows=0, body_rows=[
        ["계량(택1)", "□ 고용 실적(3명 이상)  □ 기술료 수익(1억원 이상)  □ 투자유치(협약 내 1억원 이상)"],
        ["비계량(택1)", "□ 인증 획득  □ 협력계약 체결  □ 지식재산권 확보"],
    ])
    _note(doc, "※ 사업비 구성계획")
    _grid_table(doc, ["합계(총사업비)(100%)", "정부지원금(70%)", "현금(10%)", "현물(20%)", "소계(30%)"],
                blank_rows=1)


def _static_deeptech_0_2(doc: Document) -> None:
    _heading(doc, "2. 기업 일반 현황", 13)
    _grid_table(doc, ["구분", "2025.12.31. 실적", "2026.11.30. 목표"], blank_rows=0, body_rows=[
        ["고용(명)", "", ""],
        ["매출(백만원)", "", ""],
        ["수출(백만원)", "", ""],
        ["투자(백만원)", "", ""],
    ])
    _note(doc, "※ 산업 및 지적재산권 등록현황 (신청과제 관련, 해당시)")
    _grid_table(doc, ["재산권 종류", "산업 및 지적재산권명", "등록번호(년월일)", "권리권자"], blank_rows=2)
    _note(doc, "※ 창업사업화 중복지원 검토 확인사항 (중앙정부 소관 지원사업 수행실적)")
    _grid_table(doc, ["사업명", "지원기관", "지원기간", "지원금액"], blank_rows=2)


def _static_initial_0_1(doc: Document) -> None:
    _heading(doc, "□ 일반현황", 13)
    _kv_table(doc, [
        ("기업명", ""),
        ("개업연월일", ""),
        ("사업자 구분", "□ 개인사업자  /  □ 법인사업자"),
        ("대표자 유형", "□ 단독  □ 공동  □ 각자대표"),
        ("사업자등록번호(법인등록번호)", ""),
        ("사업자 소재지(본사(점))", ""),
        ("창업아이템명", "{{ project_name }}"),
        ("산출물(협약기간 내 목표)", ""),
        ("지원 분야(택1)", "□ 제조   □ 지식서비스"),
        ("전문기술분야(택1)", "□ 기계·소재 □ 전기·전자 □ 정보·통신 □ 화공·섬유 □ 바이오·의료·생명 □ 에너지·자원 □ 공예·디자인"),
        ("지방우대 지역 해당여부", "□ 특별지원 □ 우대지원 □ 일반지역 □ 지방우대 비해당"),
    ])
    _note(doc, "※ 총 사업비 구성 계획")
    _grid_table(doc, ["정부지원사업비(A)", "자기부담 현금", "자기부담 현물", "총 사업비(C=A+B)"], blank_rows=1)
    _note(doc, "※ 팀 구성 현황 (대표자 본인 제외)")
    _grid_table(doc, ["순번", "직위", "담당 업무", "보유 역량(경력 및 학력 등)", "구성 상태"], blank_rows=2)


def _static_voucher_1(doc: Document) -> None:
    _heading(doc, "1. 혁신바우처 사업 추진 총괄책임자", 13)
    _kv_table(doc, [
        ("업체명", ""),
        ("대표자명", ""),
        ("책임자명", ""),
        ("직위", ""),
        ("휴대전화", ""),
        ("최종학력 및 전공", ""),
        ("생년월일", ""),
        ("이메일", ""),
    ])
    _note(doc, "※ 경력")
    _grid_table(doc, ["연도", "기관명", "직위", "비고"], blank_rows=2)
    _note(doc, "※ 기타 특기사항: ◦ 수상경력  ◦ 특허 출원 및 등록 등")
    doc.add_paragraph()


_STATIC_BUILDERS = {
    ("deeptech_academy", "0-1"): _static_deeptech_0_1,
    ("deeptech_academy", "0-2"): _static_deeptech_0_2,
    ("initial_package", "0-1"): _static_initial_0_1,
    ("innovation_voucher", "1"): _static_voucher_1,
}


def build_template(program_code: str) -> Path:
    form = load_form(program_code)
    doc = _base_doc()

    # 표지
    _title(doc, form.program_name)
    _subtitle(doc, "사업계획서")
    _subtitle(doc, "기업명: {{ business_name }}")
    doc.add_paragraph()

    static_ids = STATIC_SECTIONS.get(program_code, set())
    for sec in form.sections:
        if sec.id in static_ids:
            _STATIC_BUILDERS[(program_code, sec.id)](doc)
            continue
        _heading(doc, _section_label(sec.id, sec.title))
        _dyn(doc, sec.id)
        doc.add_paragraph()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"{program_code}.docx"
    doc.save(str(out))
    return out


def main() -> None:
    for code in ("deeptech_academy", "initial_package", "innovation_voucher"):
        out = build_template(code)
        print(f"built: {out} ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
