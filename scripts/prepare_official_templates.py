"""공식 양식 원본 DOCX → docxtpl 템플릿 가공기.

SAMPLE/ 의 공식 원본(한컴 변환본 포함)을 읽어, 파란 안내문구·예시를 제거하고
내용 주입 Jinja 태그를 심은 템플릿을 data/templates/{code}.docx 로 생성한다.
(기존 scripts/build_form_templates.py 의 '재현본 생성'을 대체 — 원본 레이아웃 100% 유지)

원본 파일:
  SAMPLE/deeptech_official.docx   (한컴 OOXML 변환)
  SAMPLE/voucher_official.docx    (한컴 OOXML 변환)
  SAMPLE/(별첨1) 2026년도 초기창업패키지(일반형) 사업계획서 양식.docx

주입 태그 (렌더러 core/docx_export.export_to_official_docx 와 계약):
  {{ project_name }}            과제명/창업아이템명 셀
  {{p sections["<id>"] }}       본문 섹션 서브독
  {{ ov["<라벨>"] }}            개요(요약) 표 셀 (딥테크 0-3, 초기창업 0-2)
  {{ kv2/kv4/g/kpi ... }}       혁신바우처 표 셀
  {{p sec3 }} / {{p sec6 }}     혁신바우처 서술 영역

실행: python scripts/prepare_official_templates.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
SAMPLE = ROOT / "SAMPLE"
OUT = ROOT / "data" / "templates"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


# ── 공통 헬퍼 ─────────────────────────────────────────────────────────

def body_items(doc):
    for el in doc.element.body:
        if el.tag.endswith("}p"):
            yield Paragraph(el, doc)
        elif el.tag.endswith("}tbl"):
            yield Table(el, doc)


def remove_el(obj) -> None:
    el = obj._element if hasattr(obj, "_element") else obj
    el.getparent().remove(el)


def new_tag_para(text: str):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_after(ref_obj, text: str) -> None:
    ref = ref_obj._element if hasattr(ref_obj, "_element") else ref_obj
    ref.addnext(new_tag_para(text))


def set_cell(cell, text: str) -> None:
    """셀 내용을 단일 문단 텍스트로 교체 (중첩 표 없는 셀 전용)."""
    for p in cell.paragraphs[1:]:
        remove_el(p)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        remove_el(r)
    if text:
        p0.add_run(text)


def set_first_para(cell, text: str) -> None:
    """셀의 첫 문단 텍스트만 교체 — 중첩 표는 보존."""
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        remove_el(r)
    p0.add_run(text)


def unique_cells(table):
    seen = set()
    for row in table.rows:
        for c in row.cells:
            key = id(c._tc)
            if key in seen:
                continue
            seen.add(key)
            yield c


def txbx_text(para) -> str:
    ts = para._element.findall(".//w:txbxContent//w:t", NS)
    return "".join(t.text or "" for t in ts)


def delete_between(doc, start_el, end_el) -> int:
    """body에서 start(미포함)~end(미포함) 사이 요소 삭제. end_el=None이면 끝까지."""
    n = 0
    started = False
    for el in list(doc.element.body):
        if el is start_el:
            started = True
            continue
        if end_el is not None and el is end_el:
            break
        if started and (el.tag.endswith("}p") or el.tag.endswith("}tbl")):
            el.getparent().remove(el)
            n += 1
    return n


# ── 혁신바우처 ────────────────────────────────────────────────────────

def prep_voucher() -> Path:
    doc = Document(str(SAMPLE / "voucher_official.docx"))
    T = doc.tables
    # ① 총괄책임자(T2): 예시 텍스트만 제거
    for c in unique_cells(T[2]):
        if "(예시)" in c.text:
            set_cell(c, "")
    # ② 생산제품(T3): 값 셀 → kv2 태그
    for row in T[3].rows:
        label = row.cells[0].text.strip().split("\n")[0].strip()
        set_cell(row.cells[1], f'{{{{ kv2["{label}"] }}}}')
    # ③ 추진목표(T4): 박스 셀 → 서술 서브독
    set_cell(T[4].rows[0].cells[0], "{{p sec3 }}")
    # ④ 신청내용(T5): 분야별 값 셀 → kv4 태그
    for row in T[5].rows:
        label = row.cells[0].text.strip()
        set_cell(row.cells[1], f'{{{{ kv4["{label}"] }}}}')
    # ⑤ 추진일정(T6): 간트 셀 → g[분야][월] 태그
    months = [c.text.strip() for c in T[6].rows[1].cells]
    for ri in range(2, len(T[6].rows)):
        field = T[6].rows[ri].cells[0].text.strip()
        for ci in range(1, len(T[6].rows[ri].cells)):
            m = months[ci] if ci < len(months) else ""
            if m.isdigit():
                set_cell(T[6].rows[ri].cells[ci], f'{{{{ g["{field}"]["{m}"] }}}}')
    # ⑥ 지원효과(T7): 첫 문단 → 정성 서술 서브독, 중첩 KPI 표 값 셀 → kpi 태그
    c00 = T[7].rows[0].cells[0]
    set_first_para(c00, "{{p sec6 }}")
    kpi_keys = iter(["고용 증가율(%)", "신규 고용인원수(명)", "매출액 증가율(%)", "매출 증가액(천원)"])
    if c00.tables:
        nt = c00.tables[0]
        for c in unique_cells(nt):
            if c.text.strip() in ("(%)", "(명)", "(천원)"):
                try:
                    set_cell(c, f'{{{{ kpi["{next(kpi_keys)}"] }}}}')
                except StopIteration:
                    break
    out = OUT / "innovation_voucher.docx"
    doc.save(str(out))
    return out


# ── 딥테크창업사관학교 ────────────────────────────────────────────────

_DT_HEADING = re.compile(r"^(1-[1-4]|2-[1-5]|3-[1-4]|4-[1-2])\.\s")
_DT_PART_HEADERS = ("1. 문제인식", "2. 실현가능성", "3. 성장전략", "4. 기업 구성")
# T1/T2 정리: 예시·안내 텍스트가 든 셀은 공란으로
_DT_CLEAR_MARKERS = (
    "한글로 기재", "OO도 OO시", "사업자등록번호 기재", "법인등록번호 기재",
    "0000. 00. 00", "00백만원", "0명 (대표자 제외)",
)


def prep_deeptech() -> Path:
    doc = Document(str(SAMPLE / "deeptech_official.docx"))
    T = doc.tables
    # 신청자 기본정보(T1)·기업 일반현황(T2): 예시 셀 공란화 + 과제명 태그
    for ti in (1, 2):
        for c in unique_cells(T[ti]):
            tx = c.text.strip()
            if "입교 후 진행하고자" in tx:
                set_cell(c, "{{ project_name }}")
            elif any(m in tx for m in _DT_CLEAR_MARKERS):
                set_cell(c, "")
    # 개요(요약) 표(T3): 라벨 다음 고유 셀 → ov 태그
    ov_map = {
        "명     칭": "명칭", "범     주": "범주", "소     개": "소개",
        "경쟁사 대비 차별성": "경쟁사 대비 차별성", "현황 및 구체화 방안": "현황 및 구체화 방안",
    }
    cells = list(unique_cells(T[3]))
    for i, c in enumerate(cells):
        tx = c.text.strip()
        if tx in ov_map and i + 1 < len(cells):
            set_cell(cells[i + 1], f'{{{{ ov["{ov_map[tx]}"] }}}}')
        elif tx == "Https://":
            set_cell(c, '{{ ov["회사 사이트"] }}')
        elif tx == "- 국내":
            set_cell(c, '{{ ov["진출 목표시장"] }}')
    # 과제명 표(T5)
    for c in unique_cells(T[5]):
        if "입교 후 진행하고자" in c.text:
            set_cell(c, "{{ project_name }}")

    # 본문 섹션: 제목 문단 유지 + 사이 전부 삭제 + 태그 삽입
    items = list(body_items(doc))
    anchors = []  # (element, sec_id or None)
    for it in items:
        if isinstance(it, Paragraph):
            m = _DT_HEADING.match(it.text.strip())
            if m:
                anchors.append((it._element, m.group(1)))
        elif isinstance(it, Table):
            t0 = it.rows[0].cells[0].text.strip()
            if any(t0.startswith(h) for h in _DT_PART_HEADERS):
                anchors.append((it._element, None))
    anchors.append((None, None))  # 문서 끝
    for (el, sec), (next_el, _n) in zip(anchors, anchors[1:]):
        if sec is None or el is None:
            continue
        delete_between(doc, el, next_el)
        insert_after(el, f'{{{{p sections["{sec}"] }}}}')

    out = OUT / "deeptech_academy.docx"
    doc.save(str(out))
    return out


# ── 초기창업패키지 ────────────────────────────────────────────────────

_IP_HEADINGS = ("1. 문제 인식", "2. 실현 가능성", "3. 성장전략", "4. 팀 구성")


def prep_initial() -> Path:
    doc = Document(str(SAMPLE / "(별첨1) 2026년도 초기창업패키지(일반형) 사업계획서 양식.docx"))
    items = list(body_items(doc))

    # 표지 텍스트박스 문단(사업계획서 제목) 찾기 → 그 앞(목차 페이지) 전부 삭제
    cover_el = None
    for it in items:
        if isinstance(it, Paragraph):
            tb = txbx_text(it)
            if "사업계획서" in tb and "목차" not in tb:
                cover_el = it._element
                break
    if cover_el is not None:
        for el in list(doc.element.body):
            if el is cover_el:
                break
            if el.tag.endswith("}p") or el.tag.endswith("}tbl"):
                el.getparent().remove(el)

    # 작성 안내 텍스트박스(파란 안내) 삭제
    for it in list(body_items(doc)):
        if isinstance(it, Paragraph) and txbx_text(it).startswith("※ 사업계획서는"):
            remove_el(it)
            break

    T = doc.tables
    # 일반현황(T0=기본 3x4, T1=16x20): 예시 셀 공란화 + 아이템명 태그
    for ti in (0, 1):
        for c in unique_cells(T[ti]):
            tx = c.text.strip()
            if not tx:
                continue
            if "OO기술이 적용된" in tx:
                set_cell(c, "{{ project_name }}")
            elif tx.startswith("※") or "OOOOO" in tx or "OO.OO.OO" in tx \
                    or "어플리케이션(0개)" in tx or tx.startswith("OO도"):
                set_cell(c, "")
    # 개요(요약) 표(T2): ov 태그
    ov_map = {
        "명     칭": "명칭", "범     주": "범주", "아이템 개요": "아이템 개요",
        "문제 인식": "문제 인식", "실현 가능성": "실현 가능성",
        "성장전략": "성장전략", "팀 구성": "팀 구성",
    }
    cells = list(unique_cells(T[2]))
    for i, c in enumerate(cells):
        tx = c.text.strip().split("\n")[0].strip()
        key = next((v for k, v in ov_map.items() if tx.startswith(k)), None)
        if key and i + 1 < len(cells):
            set_cell(cells[i + 1], f'{{{{ ov["{key}"] }}}}')

    # 본문 4개 섹션: 제목 텍스트박스 문단 유지, 사이 전부 삭제, 태그 삽입
    items = list(body_items(doc))
    anchors = []
    for it in items:
        if isinstance(it, Paragraph):
            tb = txbx_text(it)
            for idx, h in enumerate(_IP_HEADINGS, start=1):
                if tb.startswith(h):
                    anchors.append((it._element, str(idx)))
                    break
    anchors.append((None, None))
    for (el, sec), (next_el, _n) in zip(anchors, anchors[1:]):
        if el is None:
            continue
        delete_between(doc, el, next_el)
        insert_after(el, f'{{{{p sections["{sec}"] }}}}')

    out = OUT / "initial_package.docx"
    doc.save(str(out))
    return out


def strip_custom_props(path: Path) -> None:
    """docProps/custom.xml 제거 — 한컴 저장 시 삽입되는 Fasoo 추적 속성이
    255자를 초과해 python-docx 저장을 깨뜨리므로 통째로 걷어낸다."""
    import io
    import zipfile

    src = path.read_bytes()
    buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(src)) as zin, \
            zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "docProps/custom.xml":
                continue
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                # 한컴 저장 시 dc:description에 Fasoo 추적 blob(255자 초과)이 들어가
                # python-docx 저장을 깨뜨림 → 내용 비움
                data = re.sub(rb"<dc:description>.*?</dc:description>",
                              b"<dc:description/>", data, flags=re.S)
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    b'<Override PartName="/docProps/custom.xml" '
                    b'ContentType="application/vnd.openxmlformats-officedocument.custom-properties+xml"/>',
                    b"")
            elif item.filename == "_rels/.rels":
                data = re.sub(rb'<Relationship[^>]*Target="docProps/custom.xml"[^>]*/>', b"", data)
            zout.writestr(item, data)
    path.write_bytes(buf.getvalue())


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for fn in (prep_voucher, prep_deeptech, prep_initial):
        out = fn()
        strip_custom_props(out)
        print(f"built: {out.name} ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
