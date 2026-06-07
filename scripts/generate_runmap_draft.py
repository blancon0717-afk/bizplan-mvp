"""런맵(RunMap) 사업계획서 초안 DOCX 생성 스크립트.

실행 방법:
    cd bizplan-mvp
    python scripts/generate_runmap_draft.py

출력: scripts/런맵_사업계획서_초안.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ──────────────────────────────────────────
# 스타일 상수
# ──────────────────────────────────────────
_FONT = "맑은 고딕"
_COLOR_BLACK = RGBColor(17, 17, 17)
_MD_TABLE_ROW = re.compile(r"^\s*\|")
_BULLET_LINE = re.compile(r"^\s*-\s")
_ARROW_LINE = re.compile(r"^▶")
_H1_LINE = re.compile(r"^## [0-9]")
_H2_LINE = re.compile(r"^### [0-9]")
_H3_LINE = re.compile(r"^#### ")


def _set_font(run, size_pt: float, bold: bool, color: RGBColor) -> None:
    run.font.name = _FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), _FONT)
    rfonts.set(qn("w:hAnsi"), _FONT)
    rfonts.set(qn("w:eastAsia"), _FONT)


def _set_cell_bg(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _add_runs(para, text: str, size_pt: float, bold: bool, color: RGBColor) -> None:
    if not text:
        return
    run = para.add_run(text)
    _set_font(run, size_pt, bold, color)


def _parse_md_table(lines: list[str]) -> tuple[list[list[str]], list[list[str]]]:
    def parse_row(line: str) -> list[str]:
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        return cells

    rows = [parse_row(line) for line in lines]
    sep_idx = next(
        (i for i, r in enumerate(rows) if r and all(re.match(r"^[-:\s]+$", c) for c in r)),
        -1,
    )
    headers = rows[:sep_idx] if sep_idx >= 0 else []
    body = rows[sep_idx + 1 :] if sep_idx >= 0 else rows
    return headers, body


def _render_lines(doc: Document, lines: list[str]) -> None:
    i = 0
    while i < len(lines):
        line = lines[i]

        # 마크다운 표
        if _MD_TABLE_ROW.match(line):
            table_block: list[str] = []
            while i < len(lines) and _MD_TABLE_ROW.match(lines[i]):
                table_block.append(lines[i])
                i += 1
            headers, body = _parse_md_table(table_block)
            all_rows = headers + body
            if not all_rows:
                continue
            num_cols = max(len(r) for r in all_rows)
            tbl = doc.add_table(rows=len(all_rows), cols=num_cols)
            tbl.style = "Table Grid"
            for ri, row_cells in enumerate(all_rows):
                is_header = ri < len(headers)
                for ci in range(num_cols):
                    cell = tbl.cell(ri, ci)
                    ct = (row_cells[ci] if ci < len(row_cells) else "").strip()
                    para = cell.paragraphs[0]
                    if is_header:
                        _set_cell_bg(cell, "F8FAFC")
                        _add_runs(para, ct, 9, True, _COLOR_BLACK)
                    else:
                        _add_runs(para, ct, 9, False, _COLOR_BLACK)
            p = doc.add_paragraph()
            _set_font(p.add_run(""), 4, False, _COLOR_BLACK)
            continue

        # 빈 줄
        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue

        # ## 대섹션 제목
        if _H1_LINE.match(line):
            text = line.lstrip("#").strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
            _add_runs(para, text, 13, True, _COLOR_BLACK)
            i += 1
            continue

        # ### 중섹션 제목
        if _H2_LINE.match(line):
            text = line.lstrip("#").strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(2)
            _add_runs(para, text, 11, True, _COLOR_BLACK)
            i += 1
            continue

        # #### 소소섹션 제목
        if _H3_LINE.match(line):
            text = line.lstrip("#").strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(2)
            _add_runs(para, text, 10, True, _COLOR_BLACK)
            i += 1
            continue

        # ▶ 소제목 (DRAFT_WRITING_GUIDE 핵심 서식)
        if _ARROW_LINE.match(line):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(2)
            _add_runs(para, line.strip(), 10, True, _COLOR_BLACK)
            i += 1
            continue

        # - 세부항목
        if _BULLET_LINE.match(line):
            bullet_lines: list[str] = []
            while i < len(lines) and _BULLET_LINE.match(lines[i]):
                bullet_lines.append(lines[i])
                i += 1
            for bl in bullet_lines:
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.left_indent = Cm(0.5)
                _add_runs(para, bl.strip(), 10, False, _COLOR_BLACK)
            continue

        # 일반 단락
        text_lines: list[str] = []
        while (
            i < len(lines)
            and not _MD_TABLE_ROW.match(lines[i])
            and lines[i].strip() != ""
            and not _H1_LINE.match(lines[i])
            and not _H2_LINE.match(lines[i])
            and not _H3_LINE.match(lines[i])
            and not _ARROW_LINE.match(lines[i])
            and not _BULLET_LINE.match(lines[i])
        ):
            text_lines.append(lines[i])
            i += 1
        joined = " ".join(text_lines).strip()
        if joined:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(4)
            _add_runs(para, joined, 10, False, _COLOR_BLACK)


# ──────────────────────────────────────────
# 초안 본문 (DRAFT_WRITING_GUIDE 서식 준수)
# ──────────────────────────────────────────
DRAFT_CONTENT = """\
## 1. 개발 동기 및 현황

### 1-1. 외적 동기

▶ 러닝 시작자 10명 중 7명, 3개월을 버티지 못하고 포기
- 국내 러닝 인구는 약 2,000만 명(삼성헬스 헤비유저 기준, 2024)에 달하지만, 러닝을 시작하는 사람의 70% 이상이 3개월 내 중도 포기하는 것으로 나타남. 운동 지속성 부족은 단순한 의지 문제가 아닌 구조적 동기부여 시스템의 부재에서 기인함.
- 기존 러닝 앱(Strava, 런데이, 나이키런클럽)은 거리·속도·칼로리 등 개인 기록 측정에만 집중할 뿐, 사회적 경쟁·소속감·반복 달성 욕구 등 인간의 내재적 동기를 자극하는 메커니즘이 전무함. 이로 인해 장기 사용률이 구조적으로 낮을 수밖에 없음.

▶ 게이미피케이션 + 소셜 요소 결합 시, 운동 앱 리텐션 3~4배 향상 사례 입증
- 글로벌 게이미피케이션 시장은 2023년 약 15조 원 규모이며, 피트니스·헬스케어 분야에 접목 시 사용자 리텐션이 평균 3~4배 향상된다는 복수의 연구 결과가 존재함(Gartner, 2023).
- 국내 러닝 인구의 급속한 증가와 달리 이를 지속적으로 붙들어둘 수 있는 서비스는 아직 공백 상태임. 러닝의 '지속성 부족'이라는 문제는 시장이 명확히 존재하되 제대로 된 해결책이 없는 전형적인 언더서브드 마켓임.

### 1-2. 내적 동기

▶ 러닝 크루 1만 명을 운영하며 직접 체감한 '3개월의 벽' — 런맵 개발의 출발점
- 대표자는 10년간 러닝 크루 모임장으로 활동하며 누적 회원 1만 명 이상을 직접 관리함. 이 과정에서 처음 크루에 참여한 회원 중 상당수가 3개월을 전후해 이탈하는 반복적인 패턴을 현장에서 목격함.
- 운동과학 전공 이후 피트니스 스타트업에서 3년간 서비스 기획을 담당하며 축적한 전공 지식과 커뮤니티 운영 노하우를 결합, '경쟁 심리와 소속감을 동시에 자극하는 러닝 앱'이라는 솔루션에 도달하게 됨.
- 현재 베타 테스터 100명 대상 사전 수요 조사를 완료하였으며, 대한러닝크루협회(소속 러너 약 10만 명)와 MOU 추진 중으로 초기 채널 확보 단계에 진입함.

### 1-3. 필요성

▶ 기존 앱은 '기록'만 보여줄 뿐, 내일 다시 달리게 만들지 못함
- Strava, 런데이, 나이키런클럽 등 현재 시장 주요 앱은 모두 '개인 기록 분석'이라는 동일한 접근법을 채택하고 있으며, 다른 러너와의 실시간 경쟁이나 소속감 기반의 지속 동기 메커니즘이 부재함.
- 러닝 지속성을 높이기 위해서는 단순 기록을 넘어선 '잃으면 아까운 자산(영역)'과 '함께하는 사람들(크루)'이라는 두 가지 심리적 장치가 필요함. 이는 기존 앱 어느 것도 제공하지 못하는 기능으로, 명확한 시장 공백이 존재함.

## 2. 실현가능성

### 2-1. 시장 분석

#### 2-1-1. 시장 규모 및 성장성

▶ 2,000만 러너가 만드는 3,200억 원 시장 — 국내 러닝 앱 시장 연평균 성장 중
- 국내 러닝 인구는 약 2,000만 명(삼성헬스 헤비유저 기준, 2024)이며, 이를 기반으로 한 러닝 관련 시장 규모는 약 3,200억 원으로 추산됨(한국스포츠산업협회, 2023).
- MZ세대를 중심으로 러닝 크루 문화가 빠르게 확산되고 있으며, 서울·수도권을 중심으로 러닝 대회 참가자 수와 크루 수 모두 매년 증가 추세임. 운동의 '사회적 경험화'라는 트렌드가 러닝 앱 시장의 추가 성장 동력으로 작용할 것으로 예상됨.

#### 2-1-2. TAM–SAM–SOM

| 구분 | 규모 | 산출 근거 |
|------|------|-----------|
| TAM | 약 3,200억 원 (러너 2,000만 명) | 삼성헬스 헤비유저 기준, 한국스포츠산업협회 2023 |
| SAM | 약 1,280억 원 (스마트폰 러닝 앱 사용 가능 인구 약 800만 명) | TAM 대비 스마트폰 보유 20~40대 러너 비율 추산 |
| SOM | 약 4.2억 원 (협약기간 내 목표) | 대한러닝크루협회 채널 10만 명 × 전환율 5% = 5,000명 × ARPU 7,000원/월 × 12개월 (추정) |

#### 2-1-3. 고객 특성

▶ 1차 타깃: 디지털·경쟁에 친숙한 2030세대 직장인 러너
- 스마트폰을 활용해 러닝을 즐기는 20~40대 전체를 타깃으로 하되, 1차 타깃은 디지털 친화적이고 경쟁·소속감에 강하게 반응하는 2030세대 직장인임. 이들은 게임·SNS 사용에 익숙해 영역 점령 메커니즘과 크루 랭킹 기능에 높은 반응도를 보일 것으로 예상됨.
- 대한러닝크루협회(소속 러너 약 10만 명)와 MOU 추진 중으로, 초기 유입 채널로서 검증된 수요 기반이 이미 확보된 상태임. 베타 테스터 100명 대상 사전 수요 조사도 완료함.

### 2-2. 아이템 기술 및 고도화 방안

#### 아이템 정의

▶ 달리면 땅이 생긴다 — GPS 영역 점령 기반 게이미피케이션 러닝 앱 '런맵(RunMap)'
- 사용자가 달리기를 시작하면 GPS가 자동 작동하여 뛴 경로가 지도 위에 '내 영역'으로 표기되는 iOS·Android 모바일 앱.
- 영역 지도 + 러닝 크루 매칭 + 커뮤니티(채팅·챌린지·랭킹) 3가지 핵심 기능을 결합하여 운동 지속성과 사회적 연결감을 동시에 제공함.

#### 사용 프로세스

▶ 앱 실행 → 러닝 시작 → GPS 자동 작동·경로 기록 → 지도에 내 영역 표기 → 크루 랭킹·커뮤니티 확인 → 영역 확장 또는 방어를 위해 내일 다시 달리기

#### 핵심 기술

▶ 기술 ①: GPS 기반 실시간 영역 생성 엔진
- [입력] 사용자 달리기 경로를 GPS로 실시간 수집 → [처리] 좌표 데이터를 폴리곤으로 변환, 기존 다른 러너 영역과 충돌·중첩 판정 → [결과] 내 영역이 지도 위에 시각화되고, 경쟁 러너의 영역과 경계가 충돌할 경우 알림 발생.
- '내 구역을 빼앗긴다'는 손실 회피 심리를 자극해 매일 달리고 싶은 내재적 동기를 유발함. 기존 러닝 앱의 단순 기록 표시 대비 리텐션 유발 구조가 근본적으로 다름.

▶ 기술 ②: 위치·페이스 기반 크루 자동 매칭 알고리즘
- [입력] 사용자의 러닝 위치, 평균 페이스(분/km), 선호 거리 등 데이터 수집 → [처리] 유사 페이스·위치 기반 클러스터링으로 최적 크루 추천 → [결과] 같은 동네, 비슷한 실력의 러너들과 자동 매칭, 크루 채팅·챌린지 참여 가능.
- 소속감과 함께 달리는 경험을 제공해 장기 사용률을 높이는 핵심 사회적 장치로 기능함.

#### 경쟁사 대비 차별화 포인트

| 구분 | 런맵 | Strava | 런데이 | 나이키런클럽 |
|------|------|--------|--------|------------|
| GPS 영역 점령 게이미피케이션 | O (국내 최초) | X | X | X |
| 위치·페이스 기반 크루 자동 매칭 | O | △(일부) | X | X |
| 한국 로컬 커뮤니티 특화 | O | X | △ | X |
| 30일 리텐션 목표 | 45% | ~12% | ~12% | ~12% |

#### 개발 방식

▶ 자체 개발 3인 체제 — 협약기간 내 iOS 정식 출시 목표
- iOS 개발자 1명(전 카카오 4년 경력) 자체 개발로 iOS 버전 우선 출시, Android 버전 후속 개발.
- 대표자(기획·영업), iOS 개발자(개발), 디자이너(UX) 3인 체제로 협약기간 내 베타 → 정식 출시 완료 예정.

### 2-3. 추진성과

▶ 베타 테스터 100명 수요 조사 완료 + 대한러닝크루협회 MOU 추진 중
- 베타 테스터 100명 대상 사전 수요 조사 완료. 핵심 기능(영역 지도·크루 매칭)에 대한 실수요 존재 확인.
- 대한러닝크루협회(소속 러너 약 10만 명)와 MOU 추진 중으로, 초기 대규모 유입 채널 확보 단계에 진입함.
- 2024 서울창업허브 우수 창업자 선정으로 대표자의 사업 역량 외부 검증 완료.

## 3. 성장 전략

### 3-1. 추진 전략

#### 비즈니스 모델

▶ Freemium + 구독 + 인앱결제 — 참여할수록 확장되는 과금 구조
- 무료: 기본 영역 지도·크루 매칭 무제한 제공 (신규 유입 및 바이럴 확산용)
- 유료 구독: 월 3,900원 (국내 피트니스 앱 평균 3,300~5,500원 벤치마킹 중간값) — 영역 확장 가속, 크루 우선 매칭, 고급 통계
- 인앱결제: 포션·부스터 아이템 500~2,000원 — 영역 강화·방어 아이템
- BEP 산출: 월 운영비 약 800만 원 ÷ 구독 마진 약 3,000원 = 월 구독자 약 2,700명 확보 시 손익분기. 협약기간 말 MAU 5,000명 × 전환율 10% = 500명 → 1~2년차 내 BEP 달성 목표.

#### 5개년 매출 목표

| 구분 | 1년차 | 2년차 | 3년차 | 4년차 | 5년차 |
|------|-------|-------|-------|-------|-------|
| MAU | 5,000 | 3만 | 10만 | 25만 | 50만 |
| 유료 전환율 | 10% | 15% | 20% | 22% | 25% |
| ARPU(월) | 3,900원 | 4,200원 | 4,500원 | 4,800원 | 5,000원 |
| 연 매출 목표 | 약 2,340만 원 | 약 2.3억 원 | 약 10.8억 원 | 약 31.7억 원 | 약 75억 원 |

#### 마케팅 전략

▶ 커뮤니티 → 바이럴 → 광고 3단계 진입
- 1단계 (협약기간): 대한러닝크루협회 채널 직접 유입 + 대표자 러닝 크루 네트워크(누적 1만 명) 활용한 입소문 확산. 베타 테스터 100명을 '앰배서더'로 전환, 리뷰·챌린지 콘텐츠 생산 유도.
- 2단계 (2년차): 영역 점령 지도 이미지 공유 → SNS 자연 바이럴 유도. '내 동네를 내가 점령했다'는 시각적 달성감은 그 자체로 강력한 공유 콘텐츠로 기능함.
- 3단계 (3년차~): 성과 데이터 기반 유료 광고 집행(유튜브·인스타그램), 기업 복지몰 및 스포츠 브랜드 파트너십 제휴.

#### 사업 전체 로드맵

| 순번 | 추진 내용 | 추진 기간 | 세부 내용 |
|------|-----------|-----------|-----------|
| 1 | iOS 베타 출시 | 협약 4개월차 | 핵심 3기능(영역·크루·커뮤니티) 탑재, 베타 테스터 300명 운영 |
| 2 | 정식 출시 + MAU 5,000 달성 | 협약 종료 | App Store 정식 등록, 대한러닝크루협회 채널 오픈 |
| 3 | Android 출시 + 구독 BEP 달성 | 2년차 | Android 버전 출시, MAU 3만 목표 |
| 4 | 해외 진출 (일본·동남아) | 3년차~ | 영역 게이미피케이션 특화 일본 러닝 커뮤니티 우선 진입 |

#### 기술 보호 계획

▶ GPS 영역 알고리즘 특허 출원 계획 (협약기간 내)
- GPS 기반 영역 생성 알고리즘 및 크루 자동 매칭 알고리즘에 대한 특허 출원 예정 (협약기간 내).

### 3-2. 자금 계획

#### 사업비 집행 계획

| 비 목 | 집행 계획 | 정부지원사업비(원) | 자기부담사업비(현금) | 자기부담사업비(현물) | 합계 |
|-------|-----------|-------------------|---------------------|---------------------|------|
| 인건비 | 대표 + iOS 개발자 + 디자이너 (3인, 협약기간) | 60,000,000 | — | 20,000,000 | 80,000,000 |
| 외주용역비 | Android 개발 외주, QA 테스트 | 15,000,000 | 5,000,000 | — | 20,000,000 |
| 지급수수료 | App Store·Google Play 개발자 등록, 법률 자문 | 3,000,000 | — | — | 3,000,000 |
| 광고선전비 | 베타 테스터 모집, 출시 이벤트 | 12,000,000 | 3,000,000 | — | 15,000,000 |
| 특허권 등 무형자산취득 | GPS 영역 알고리즘 특허 출원 | 10,000,000 | — | — | 10,000,000 |
| 합 계 | | 100,000,000 | 8,000,000 | 20,000,000 | 128,000,000 |

#### 자금 조달 계획

| 구분 | 시기 | 조달 금액 | 조달 방안 | 자금 용도 |
|------|------|-----------|-----------|-----------|
| 1단계 | 협약기간 | 1.28억 원 | 정부지원금 + 자기부담 | iOS 출시, 베타 운영, 특허 출원 |
| 2단계 | 2년차 | 3~5억 원 | Seed 투자유치 + 정책자금 | Android 출시, MAU 3만 확보, 마케팅 |
| 3단계 | 3년차~ | 10~20억 원 | 시리즈 A + 자체 영업이익 재투자 | 해외 진출(일본·동남아), 기술 고도화 |

## 4. 기업 구성

### 4-1. 기업 구성

#### 대표자 현황 및 역량

▶ 러닝 크루 1만 명 운영 + 운동과학 전공 + 스타트업 기획 3년 — 이 문제를 해결할 가장 적합한 창업자
- 대표자는 운동과학 전공 후 피트니스 스타트업에서 3년간 서비스 기획을 담당하며 앱 서비스 개발 프로세스와 스포츠 분야 사용자 니즈를 동시에 체득함.
- 10년간 러닝 크루 모임장으로 활동하며 누적 회원 1만 명 이상을 운영한 커뮤니티 운영 경험은 '러닝 지속성 부족'이라는 문제를 가장 깊이 이해하고 있는 배경이 됨.
- 2024 서울창업허브 우수 창업자 선정으로 외부 검증 완료.

#### 팀원 및 고용 계획

| 구성원 | 역할 | 주요 역량 |
|--------|------|-----------|
| 대표자 | 기획·영업·마케팅 | 운동과학 전공, 피트니스 스타트업 기획 3년, 러닝 크루 운영 10년 |
| iOS 개발자 | iOS 앱 개발 전담 | 전 카카오 4년 경력, iOS 서비스 개발 실무 검증 |
| 디자이너 | UX/UI 디자인 | UX 전문, 모바일 앱 인터페이스 설계 역량 |

▶ 3인 전원 4대보험 가입 예정 (지원금 수령 후 즉시 등록)
- 2년차 Android 개발자 1명 신규 채용 계획으로 일자리 창출 의지 보유.
"""


def build_docx() -> None:
    doc = Document()

    # 여백 설정 (25mm)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Normal 스타일 기본 폰트 (맑은 고딕)
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(10)
    style_rpr = style.element.find(qn("w:rPr"))
    if style_rpr is not None:
        style_rfonts = style_rpr.find(qn("w:rFonts"))
        if style_rfonts is not None:
            style_rfonts.set(qn("w:eastAsia"), _FONT)

    # 표지
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    _add_runs(title_para, "런맵(RunMap) 사업계획서 초안", 16, True, _COLOR_BLACK)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(2)
    _add_runs(sub_para, "GPS 영역 점령 기반 게이미피케이션 러닝 앱", 11, False, _COLOR_BLACK)

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(info_para, "※ 본 문서는 인터뷰 답변 기반 AI 초안이며 수치·출처 보완이 필요함.", 9, False, _COLOR_BLACK)

    doc.add_paragraph()

    # 본문 렌더링
    lines = DRAFT_CONTENT.split("\n")
    _render_lines(doc, lines)

    # 저장
    out_path = Path(__file__).parent / "런맵_사업계획서_초안.docx"
    doc.save(out_path)
    print(f"저장 완료: {out_path}")


if __name__ == "__main__":
    build_docx()
