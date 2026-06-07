"""사업계획서 초안 생성 테스트 스크립트 (백엔드 무관, 독립 실행).

사용법:
    python scripts/test_draft_gen.py claude-haiku-4-5-20251001
    python scripts/test_draft_gen.py claude-sonnet-4-6
    python scripts/test_draft_gen.py claude-opus-4-8

출력: scripts/output/사업계획서_초안_{model}.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import os
import time

from dotenv import load_dotenv

sys.path.insert(0, str(Path(__file__).parent.parent))

load_dotenv()

from anthropic import Anthropic

# ──────────────────────────────────────────
# 인터뷰 답변 (반려동물 용품 커머스 플랫폼)
# ──────────────────────────────────────────
INTERVIEW = """\
아이템: 반려동물 용품 커머스 플랫폼 (맞춤 큐레이션 앱)

Q1. 해당 사업을 하게 된 이유는 무엇이며, 대표님의 경력·이력과 어떻게 연관되나요?
A: 강아지를 키우다 보니 좋은 제품 찾기 어렵더라고요. 그래서 시작했어요. 저는 원래 회사 다녔는데 마케팅 쪽이요. 한 3년 정도?

Q2. 이 사업으로 핵심적으로 해결하고자 하는 문제는 무엇인가요? 기존에는 고객들이 어떻게 해결하고 있었고, 그 방식의 한계는 무엇인가요?
A: 기존에는 그냥 쿠팡이나 네이버에서 샀는데, 뭐가 좋은지 모르잖아요. 정보가 너무 많고. 그게 불편하니까 저희가 큐레이션 해주는 거예요.

Q3. 주 고객은 누구이며, 그 중 1차 타깃은 누구인가요? 고객 특성의 근거나 확보된 수요처가 있으면 함께 알려주세요.
A: 반려동물 키우는 사람들이요. 주로 2030 여성? 일단 주변에서 좋다고 했어요. 지인들한테 물어봤는데 다들 쓰고 싶다고.

Q4. 목표 시장 규모는 얼마이며, 어떤 기준으로 측정하셨나요?
A: 반려동물 시장이 엄청 크잖아요. 한 4조인가 그렇다고 어디서 봤어요. 계속 크고 있고요.

Q5. 협약기간 내 만들려는 아이템/제품은 무엇이며, 핵심 기능은 무엇인가요?
A: 앱이요. 반려동물 정보 입력하면 맞춤 상품 추천해주는 거. 그리고 구매도 되고요.

Q6. 그 핵심 기능이 앞서 말씀하신 문제를 어떻게 해결하며, 소비자는 어떤 효용을 얻나요?
A: 뭐 고를지 모르는 사람한테 바로 추천해주니까 편하죠. 시간도 절약되고.

Q7. 주요 경쟁사(또는 유사 서비스)는 무엇이며, 자사만의 차별적 강점은 무엇인가요?
A: 핏펫? 그런 데 있는데, 저희는 큐레이션이 더 정확해요. 알고리즘으로 하니까요.

Q8. 비즈니스 모델은 어떻게 되나요? 누구에게, 얼마에, 어떤 구조로 파시나요?
A: 수수료 모델이에요. 판매되면 저희가 일정 % 가져가는 구조요. 몇 %인지는 아직 정확히 안 정했어요.

Q9. 향후 5년 예상 매출과 그 근거는 무엇인가요?
A: 1년차에 1억, 3년차에 10억, 5년차에 100억 정도? 시장이 크니까 잘 되면 될 것 같아요.

Q10. 대표자 역량·경력은 어떻게 되며, 팀원 구성(인원·역할)은 어떻게 되나요?
A: 저 마케팅 3년 했고요, 개발자 친구 한 명 있어요. 디자인은 프리랜서 쓸 예정이에요. 일단 셋이서 해보려고요.
"""

# ──────────────────────────────────────────
# DOCX 스타일 상수
# ──────────────────────────────────────────
_FONT = "맑은 고딕"
_BLACK = RGBColor(17, 17, 17)
_MD_TABLE = re.compile(r"^\s*\|")
_BULLET = re.compile(r"^\s*-\s")
_ARROW = re.compile(r"^▶")
_H1 = re.compile(r"^## \d")
_H2 = re.compile(r"^### \d")
_H3 = re.compile(r"^#### ")


def _set_font(run, size: float, bold: bool, color: RGBColor = _BLACK) -> None:
    run.font.name = _FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), _FONT)
    rfonts.set(qn("w:hAnsi"), _FONT)
    rfonts.set(qn("w:eastAsia"), _FONT)


def _set_cell_bg(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _add_run(para, text: str, size: float, bold: bool) -> None:
    if text:
        _set_font(para.add_run(text), size, bold)


def _parse_table(lines: list[str]) -> tuple[list[list[str]], list[list[str]]]:
    def row(line: str) -> list[str]:
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        return cells

    rows = [row(l) for l in lines]
    sep = next((i for i, r in enumerate(rows) if r and all(re.match(r"^[-:\s]+$", c) for c in r)), -1)
    return rows[:sep] if sep >= 0 else [], rows[sep + 1:] if sep >= 0 else rows


def _render(doc: Document, lines: list[str]) -> None:
    i = 0
    while i < len(lines):
        line = lines[i]

        if _MD_TABLE.match(line):
            block: list[str] = []
            while i < len(lines) and _MD_TABLE.match(lines[i]):
                block.append(lines[i])
                i += 1
            headers, body = _parse_table(block)
            all_rows = headers + body
            if not all_rows:
                continue
            cols = max(len(r) for r in all_rows)
            tbl = doc.add_table(rows=len(all_rows), cols=cols)
            tbl.style = "Table Grid"
            for ri, row_cells in enumerate(all_rows):
                is_hdr = ri < len(headers)
                for ci in range(cols):
                    cell = tbl.cell(ri, ci)
                    ct = (row_cells[ci] if ci < len(row_cells) else "").strip()
                    if is_hdr:
                        _set_cell_bg(cell, "F8FAFC")
                    _add_run(cell.paragraphs[0], ct, 9, is_hdr)
            doc.add_paragraph()
            continue

        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue

        if _H1.match(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            _add_run(p, line.lstrip("#").strip(), 13, True)
            i += 1
            continue

        if _H2.match(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, line.lstrip("#").strip(), 11, True)
            i += 1
            continue

        if _H3.match(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, line.lstrip("#").strip(), 10, True)
            i += 1
            continue

        if _ARROW.match(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, line.strip(), 10, True)
            i += 1
            continue

        if _BULLET.match(line):
            while i < len(lines) and _BULLET.match(lines[i]):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(2)
                _add_run(p, lines[i].strip(), 10, False)
                i += 1
            continue

        # 일반 단락
        text_lines: list[str] = []
        while (
            i < len(lines)
            and not _MD_TABLE.match(lines[i])
            and lines[i].strip() != ""
            and not _H1.match(lines[i])
            and not _H2.match(lines[i])
            and not _H3.match(lines[i])
            and not _ARROW.match(lines[i])
            and not _BULLET.match(lines[i])
        ):
            text_lines.append(lines[i])
            i += 1
        joined = " ".join(text_lines).strip()
        if joined:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _add_run(p, joined, 10, False)


def _save_docx(content: str, model: str) -> Path:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    _add_run(title, "반려동물 용품 커머스 플랫폼 사업계획서 초안", 16, True)

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(label, f"생성 모델: {model}", 9, False)

    doc.add_paragraph()

    _render(doc, content.split("\n"))

    out_dir = Path(__file__).parent / "output"
    out_dir.mkdir(exist_ok=True)
    short = model.replace("claude-", "").replace("-20251001", "")
    out_path = out_dir / f"사업계획서_초안_{short}.docx"
    doc.save(str(out_path))
    return out_path


def _generate(model: str) -> str:
    draft_guide = (
        Path(__file__).parent.parent / "skills" / "L1_universal" / "DRAFT_WRITING_GUIDE.md"
    ).read_text(encoding="utf-8")

    system = f"""\
당신은 정부지원사업 사업계획서 전문 컨설턴트입니다.
아래 DRAFT_WRITING_GUIDE의 서식과 작성 원칙을 정확히 따라 사업계획서 초안을 작성합니다.

# DRAFT_WRITING_GUIDE
{draft_guide}

## 출력 서식 규칙
- 대섹션: `## 1. 개발 동기 및 현황` 형식
- 중섹션: `### 1-1. 외적 동기` 형식
- 소소섹션: `#### 2-2-1. 아이템 정의` 형식
- 소제목: `▶ 메시지 중심 문구` 형식
- 세부항목: `- 내용` 형식
- 표: 마크다운 표
- 문장 종결: 음슴체 (~함, ~임, ~됨)
- 수치 출처 없으면: (추정) 표기
- 인터뷰에 없는 정보는 합리적으로 추정 후 (추정) 표기, 플레이스홀더 금지"""

    user = f"""\
아래 인터뷰 답변을 바탕으로 DRAFT_WRITING_GUIDE의 4개 대섹션을 모두 작성하세요.

## 인터뷰 답변
{INTERVIEW}

## 작성할 섹션 (전부 포함 필수)
## 1. 개발 동기 및 현황
### 1-1. 외적 동기
### 1-2. 내적 동기
### 1-3. 필요성

## 2. 실현가능성
### 2-1. 시장 분석
#### 2-1-1. 시장 규모 및 성장성
#### 2-1-2. TAM-SAM-SOM
#### 2-1-3. 고객 특성
### 2-2. 아이템 기술 및 고도화 방안
#### 2-2-1. 아이템 정의
#### 2-2-2. 사용 프로세스
#### 2-2-3. 핵심 기술
#### 2-2-4. 개발 방식
#### 2-2-5. 차별화 포인트 (경쟁사 비교표 + 부연설명 3개)
### 2-3. 추진성과

## 3. 성장 전략
### 3-1. 추진 전략
#### 3-1-1. 비즈니스 모델 (BEP 산출 포함)
#### 3-1-2. 5개년 매출 목표 (표)
#### 3-1-3. 마케팅 전략
#### 3-1-4. 사업 전체 로드맵 (표)
#### 3-1-5. 사업추진일정 — 표 A(KPI), 표 B(월별 일정)
#### 3-1-6. 기술 보호 계획 (IP 표)
### 3-2. 자금 계획
#### 3-2-1. 사업비 집행 계획 (표, 합계 1억 4,300만원 기준)
#### 3-2-2. 자금 조달 계획 (표)

## 4. 기업 구성
### 4-1. 기업 구성
#### 4-1-1. 대표자 현황 및 역량
#### 4-1-2. 기업 현황 및 역량
#### 4-1-3. 조직 구성 현황 (표)
#### 4-1-4. 협약기간 내 추가 고용 계획 (표)
#### 4-1-5. 외부 네트워크 현황 (표)"""

    client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

    create_kwargs: dict = dict(
        model=model,
        max_tokens=8192,
        system=system,
        messages=[{"role": "user", "content": user}],
    )
    # Opus 4.8은 temperature 미지원
    if "opus-4-8" not in model:
        create_kwargs["temperature"] = 0.3

    start = time.time()
    resp = client.messages.create(**create_kwargs)
    elapsed = int((time.time() - start) * 1000)

    text = "".join(b.text for b in resp.content if getattr(b, "type", "") == "text")
    print(f"  입력 {resp.usage.input_tokens:,}토큰 / 출력 {resp.usage.output_tokens:,}토큰 / {elapsed:,}ms")
    return text


def main() -> None:
    model = sys.argv[1] if len(sys.argv) > 1 else "claude-haiku-4-5-20251001"
    print(f"[시작] 모델: {model}")
    content = _generate(model)
    out = _save_docx(content, model)
    print(f"[완료] {out}")


if __name__ == "__main__":
    main()
